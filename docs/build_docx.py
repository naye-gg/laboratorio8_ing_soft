"""Genera docs/ARQUITECTURA.docx a partir del contenido del documento de arquitectura.

Construye un Word con encabezados, tablas e imagenes (diagramas de Eraser)
incrustadas. Ejecutar:  .venv/bin/python docs/build_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
IMG_W = Inches(6.3)  # ancho util de pagina carta con margenes por defecto
REPO_URL = "https://github.com/naye-gg/Laboratorio8_ingsoft"
SONAR_URL = "https://sonarqube.ingsoftware.lat/dashboard?id=Nayeli_Guerrero_t1"


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Inserta un hyperlink clickable (azul, subrayado) en el parrafo dado."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    link.append(run)
    paragraph._p.append(link)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def add_image(doc: Document, name: str, caption: str) -> None:
    path = DOCS / name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=IMG_W)
    add_caption(doc, caption)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


doc = Document()

# ---- Portada / titulo ----
title = doc.add_heading("Documento de Arquitectura", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Programa de Recompensas de Restaurantes").bold = True
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Curso: CS3081 - Ingenieria de Software")
tarea = doc.add_paragraph()
tarea.alignment = WD_ALIGN_PARAGRAPH.CENTER
tarea.add_run("Tarea: Laboratorio 8")
est = doc.add_paragraph()
est.alignment = WD_ALIGN_PARAGRAPH.CENTER
est.add_run("Estudiante: Nayeli Guerrero Gutierrez - 202410790")
repo = doc.add_paragraph()
repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo.add_run("Repositorio: ")
add_hyperlink(repo, REPO_URL, REPO_URL)

# ---- 1. Resumen ----
doc.add_heading("1. Resumen", level=1)
doc.add_paragraph(
    "Sistema que transforma el consumo de un cliente en un restaurante afiliado en "
    "puntos y cashback, siguiendo una arquitectura orientada a eventos (EDA) sobre "
    "RabbitMQ y organizada internamente con Arquitectura Hexagonal (puertos y adaptadores)."
)
doc.add_paragraph("El flujo implementa exactamente el proceso de la Figura 1 del enunciado:")
add_image(doc, "figura1_enunciado.png", "Figura 1. Proceso del programa de recompensas (enunciado).")
doc.add_paragraph(
    "Reproduccion del proceso de la Figura 1 del enunciado (con RabbitMQ/AMQP en lugar "
    "de ActiveMQ): productor que registra la cena -> Publish del mensaje con la "
    "informacion del usuario y la cena -> Exchange enruta a la Queue -> el consumidor "
    "registra los puntos o el reembolso."
)
add_image(doc, "proceso_figura1_eraser.png", "Figura 2. Proceso implementado con RabbitMQ/AMQP.")

# ---- 2. Patron arquitectonico ----
doc.add_heading("2. Patron arquitectonico: Hexagonal + EDA", level=1)
doc.add_paragraph(
    "Por que Hexagonal + EDA. El enunciado exige alta cohesion, bajo acoplamiento, "
    "modularidad y escalabilidad sobre un escenario de mensajeria. Se eligio Arquitectura "
    "Orientada a Eventos (sobre RabbitMQ) porque desacopla al productor (restaurante) del "
    "consumidor (recompensas): se comunican solo por un evento en una cola, lo que permite "
    "escalar y anadir consumidores sin tocar al productor. Y se eligio Arquitectura Hexagonal "
    "para aislar la logica de negocio de la infraestructura mediante puertos y adaptadores, "
    "logrando bajo acoplamiento, alta testabilidad y la posibilidad de cambiar el broker o la "
    "persistencia sin afectar el dominio. Ambos patrones son los recomendados por el enunciado "
    "y se complementan: EDA resuelve la comunicacion; Hexagonal, la organizacion interna."
)
doc.add_paragraph(
    "La regla de dependencia es estricta: todo apunta hacia el dominio. La infraestructura "
    "depende de la aplicacion y el dominio; nunca al reves."
)
doc.add_paragraph(
    "El siguiente diagrama (Eraser) muestra la arquitectura real en la notacion de puertos "
    "y adaptadores: los driving adapters (inbound) a la izquierda — CLI del restaurante y "
    "RabbitMQEventConsumer —, el nucleo hexagonal al centro (capa de aplicacion con los "
    "casos de uso, rodeando a la capa de dominio con entidades, RewardPolicy, value objects "
    "y eventos), y los driven adapters (outbound) a la derecha — RabbitMQEventPublisher, "
    "serialization e InMemoryRewardAccountRepository — conectados a traves de los puertos. "
    "Las dependencias siempre apuntan hacia el dominio."
)
add_image(doc, "arquitectura_eraser.png", "Figura 3. Arquitectura hexagonal (puertos y adaptadores).")

doc.add_heading("Puertos (interfaces) y adaptadores (implementaciones)", level=2)
table(
    doc,
    ["Puerto (dominio)", "Adaptador (infraestructura)"],
    [
        ["EventPublisher", "RabbitMQEventPublisher"],
        ["EventConsumer", "RabbitMQEventConsumer"],
        ["RewardAccountRepository", "InMemoryRewardAccountRepository"],
    ],
)
doc.add_paragraph(
    "Gracias a esta inversion de dependencias, el dominio no conoce RabbitMQ: cambiar el "
    "broker (a Kafka o ActiveMQ) o la persistencia (a PostgreSQL) solo requiere escribir un "
    "nuevo adaptador, sin tocar la logica de negocio."
)

# ---- 3. Atributos de calidad ----
doc.add_heading("3. Como se cumplen los atributos de calidad", level=1)
for titulo, cuerpo in [
    ("Alta cohesion", "cada modulo tiene una sola responsabilidad. La regla de fidelizacion vive unicamente en RewardPolicy; la validacion de datos, en los objetos de valor (Money, CardNumber, RestaurantCode)."),
    ("Bajo acoplamiento", "las capas se comunican mediante puertos abstractos (Protocol) y eventos inmutables. Productor y consumidor solo comparten el contrato del evento, no codigo."),
    ("Modularidad", "paquetes independientes (domain, application, infrastructure, apps) sustituibles de forma aislada."),
    ("Abstraccion", "el nucleo trabaja con conceptos de negocio (cena, recompensa, cuenta), no con detalles tecnicos."),
    ("Escalabilidad", "al desacoplar mediante colas, productor y consumidor escalan de forma independiente; se pueden anadir mas consumidores sobre la misma cola."),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{titulo}: ").bold = True
    p.add_run(cuerpo)

# ---- 4. Componentes ----
doc.add_heading("4. Componentes principales", level=1)
table(
    doc,
    ["Capa", "Archivo", "Responsabilidad"],
    [
        ["Dominio", "value_objects.py", "Invariantes de Money, CardNumber, RestaurantCode"],
        ["Dominio", "model.py", "Entidades Dinner, Reward, agregado RewardAccount"],
        ["Dominio", "events.py", "Eventos DinnerRegistered, RewardProcessed"],
        ["Dominio", "reward_policy.py", "Calculo de puntos/cashback (regla de negocio)"],
        ["Dominio", "ports.py", "Contratos (puertos) de la hexagonal"],
        ["Aplicacion", "register_dinner.py", "Caso de uso del productor"],
        ["Aplicacion", "process_reward.py", "Caso de uso del consumidor"],
        ["Infraestructura", "messaging/", "Adaptadores RabbitMQ (publisher/consumer/conexion)"],
        ["Infraestructura", "persistence/", "Repositorio de cuentas"],
        ["Infraestructura", "serialization.py", "Traduccion evento <-> JSON"],
        ["Apps", "producer_app.py, consumer_app.py", "Composition root ejecutable"],
    ],
)

# ---- 5. Decisiones de diseno ----
doc.add_heading("5. Decisiones de diseno relevantes", level=1)
for titulo, cuerpo in [
    ("Seguridad", "el numero de tarjeta nunca se expone completo; CardNumber.masked revela solo los ultimos 4 digitos y el evento RewardProcessed se serializa enmascarado. Las credenciales se leen de variables de entorno y la contrasena se excluye del repr de la configuracion."),
    ("Confiabilidad", "los mensajes invalidos se descartan de forma controlada (SerializationError) sin tumbar el consumidor; las colas se declaran durable=True."),
    ("Testabilidad", "la dependencia con pika se aisla y se inyecta el canal, lo que permite probar los adaptadores con dobles de prueba. El tiempo se inyecta mediante un clock para obtener pruebas deterministas."),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{titulo}: ").bold = True
    p.add_run(cuerpo)

# ---- 6. Calidad y pruebas ----
doc.add_heading("6. Calidad y pruebas", level=1)
for item in [
    "Pruebas automatizadas con pytest (54 pruebas).",
    "Cobertura 100 % (medida con pytest-cov, reporte coverage.xml; minimo exigido 85 %).",
    "Analisis estatico y control de duplicidad con SonarQube (sonar-scanner): Reliability A, Security A, Maintainability A, Duplicaciones 0 %, 0 Security Hotspots.",
]:
    doc.add_paragraph(item, style="List Bullet")

sonar = doc.add_paragraph()
sonar.add_run("Proyecto en SonarQube: ")
add_hyperlink(sonar, SONAR_URL, SONAR_URL)
sonar.add_run(". El Quality Gate se encuentra en estado Passed:")
add_image(
    doc,
    "evidencia/sonarqube_dashboard.png",
    "SonarQube — Quality Gate \"Passed\": Coverage 100 %, 0 issues, "
    "Duplicaciones 0 %, 0 Security Hotspots (Security A).",
)

# ---- 7. Casos de uso ----
doc.add_heading("7. Diagrama de casos de uso", level=1)
add_image(doc, "casos_de_uso_definitivo.png", "Figura 4. Diagrama de casos de uso.")
doc.add_paragraph(
    "Actores: Restaurante Afiliado y Cliente (izquierda), Sistema de Notificaciones "
    "(derecha). Se modelan las relaciones <<include>>, <<extend>> y el disparo por "
    "evento (<<trigger>>) entre casos de uso."
)
doc.add_paragraph(
    "Diagramas de arquitectura y de proceso editables en Eraser (espacio \"Libre\"): "
    "https://app.eraser.io/workspace/2rJMGe6QAMfV3oAHAUuV."
)

# ---- 8. Evidencia de pruebas ----
doc.add_heading("8. Evidencia de pruebas", level=1)
for item in [
    "docs/evidencia/pytest_salida.txt: salida completa de pytest -v (54 pruebas, 100% cobertura).",
    "docs/evidencia/cobertura_html/index.html: reporte navegable de cobertura.",
    "docs/evidencia/sonarqube_dashboard.png: captura del Quality Gate (Passed) en SonarQube.",
]:
    doc.add_paragraph(item, style="List Bullet")

out = DOCS / "ARQUITECTURA.docx"
doc.save(str(out))
print(f"OK -> {out}")
